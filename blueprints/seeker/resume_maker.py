import os
from pymongo import MongoClient
from flask import Blueprint, render_template, request, jsonify, current_app, flash, redirect, url_for
from bson.objectid import ObjectId
from flask_login import login_required, current_user
from datetime import datetime
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
import re
import spacy
import io
from bson import json_util
import json
import yaml
import google.generativeai as genai

resume_maker_bp = Blueprint('resume_maker', __name__)
MONGO_URI = os.getenv("MONGO_URI", "mongodb+srv://AronJain:4E1zkxYGeaWZQCL8@cluster0.qy4jgjm.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
client = MongoClient(MONGO_URI)
db = client['job_portal']
collection_resume_details = db.tbl_resume_details

# Load spaCy's pre-trained model
nlp = spacy.load("en_core_web_sm")

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@resume_maker_bp.route('/select-template')
@login_required
def select_template():
    return render_template('seeker/select_template.html')

@resume_maker_bp.route('/setup-profile')
@login_required
def setup_profile():
    return render_template('seeker/setup_profile.html')

@resume_maker_bp.route('/resume-parser')
@login_required
def resume_parser():
    return render_template('seeker/resume_parser.html')

@resume_maker_bp.route('/resume-maker')
@login_required
def resume_maker():
    # Check if user already has a parsed resume
    existing_resume = collection_resume_details.find_one({'user_id': ObjectId(current_user.id)})
    template = request.args.get('template', 'modern')
    
    return render_template(
        'seeker/resume_maker.html',
        template=template,
        existing_data=existing_resume
    )

@resume_maker_bp.route('/save-resume', methods=['POST'])
@login_required
def save_resume():
    try:
        data = request.json
        
        resume_data = {
            'user_id': ObjectId(current_user.id),
            'personal_info': data.get('personal_info', {}),
            'professional_summary': data.get('professional_summary', ''),
            'education': data.get('education', []),
            'skills': data.get('skills', {}),
            'work_experience': data.get('work_experience', []),
            'projects': data.get('projects', []),
            'languages': data.get('languages', []),
            'achievements': data.get('achievements', []),
            'last_updated': datetime.utcnow()
        }
        
        # Update or insert resume in database
        result = collection_resume_details.update_one(
            {'user_id': ObjectId(current_user.id)},
            {'$set': resume_data},
            upsert=True
        )
        
        return jsonify({
            'status': 'success',
            'message': 'Resume saved successfully'
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@resume_maker_bp.route('/get-resume', methods=['GET'])
@login_required
def get_resume():
    try:
        # Fetch existing resume data
        resume = current_app.collection_resume_details.find_one(
            {'user_id': ObjectId(current_user.id)}
        )
        
        if resume:
            # Convert ObjectId to string for JSON serialization
            resume['_id'] = str(resume['_id'])
            resume['user_id'] = str(resume['user_id'])
            return jsonify({
                'status': 'success',
                'data': resume
            })
        else:
            return jsonify({
                'status': 'success',
                'data': None
            })
            
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@resume_maker_bp.route('/generate-pdf', methods=['POST'])
@login_required
def generate_pdf():
    try:
        resume_data = request.json
        # Add PDF generation logic here
        
        return jsonify({
            'status': 'success',
            'message': 'PDF generated successfully',
            'pdf_url': 'url_to_generated_pdf'
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500 

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file object."""
    with pdfplumber.open(pdf_file) as pdf:
        text = " ".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file object."""
    doc = Document(docx_file)
    text = " ".join([para.text for para in doc.paragraphs])
    return text

def extract_email(text):
    """Extract email using regex."""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else None

def extract_phone(text):
    """Extract phone number using regex."""
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    return phones[0] if phones else None

def extract_name(text):
    """Extract name using NLP."""
    doc = nlp(text[:1000])  # Process first 1000 chars for efficiency
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None

def extract_skills(text):
    """Extract skills from the text."""
    skills_section = re.search(r"Skills?(.*?)(Education|Experience|$)", text, re.DOTALL | re.IGNORECASE)
    if skills_section:
        skills_text = skills_section.group(1).strip()
        # Split by common delimiters
        skills = re.split(r"[,•|\n]", skills_text)
        # Clean and filter skills
        skills = [skill.strip() for skill in skills if skill.strip()]
        # Separate technical and soft skills (you might want to customize this list)
        technical_skills = []
        soft_skills = []
        for skill in skills:
            if any(tech in skill.lower() for tech in ['python', 'java', 'sql', 'html', 'css', 'javascript']):
                technical_skills.append(skill)
            else:
                soft_skills.append(skill)
        return {'technical': technical_skills, 'soft': soft_skills}
    return {'technical': [], 'soft': []}

def extract_education(text):
    """Extract education details."""
    education = []
    education_section = re.search(r"Education(.*?)(Experience|Skills|$)", text, re.DOTALL | re.IGNORECASE)
    if education_section:
        edu_text = education_section.group(1)
        # Look for degree patterns
        degree_patterns = [
            r"(Bachelor|Master|PhD|B\.|M\.|Ph\.D)\.?\s+(?:of|in)?\s+([^,\n]+)",
            r"([^,\n]+(?:University|College|Institute)[^,\n]+)"
        ]
        for pattern in degree_patterns:
            matches = re.finditer(pattern, edu_text, re.IGNORECASE)
            for match in matches:
                education.append({
                    "degree": match.group(0).strip(),
                    "institution": "",  # You might want to extract this separately
                    "graduation_year": ""  # You might want to extract this separately
                })
    return education

def extract_experience(text):
    """Extract work experience."""
    experience = []
    exp_section = re.search(r"Experience(.*?)(Education|Skills|$)", text, re.DOTALL | re.IGNORECASE)
    if exp_section:
        exp_text = exp_section.group(1)
        # Split into different positions
        positions = re.split(r"\n(?=\d{4}|January|February|March|April|May|June|July|August|September|October|November|December)", exp_text)
        for position in positions:
            if position.strip():
                # Extract company and title
                company_match = re.search(r"(?:at|@)\s*([^,\n]+)", position)
                title_match = re.search(r"([^,\n]+(?:Engineer|Developer|Manager|Director|Analyst)[^,\n]+)", position)
                experience.append({
                    "company": company_match.group(1).strip() if company_match else "",
                    "title": title_match.group(1).strip() if title_match else "",
                    "duration": "",  # You might want to extract this
                    "responsibilities": [position.strip()]  # You might want to split this further
                })
    return experience

@resume_maker_bp.route('/parse-resume', methods=['POST'])
@login_required
def parse_resume():
    try:
        if 'file' not in request.files:
            current_app.logger.error('No file in request')
            return jsonify({'status': 'error', 'message': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            current_app.logger.error('No filename')
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400

        if file and allowed_file(file.filename):
            try:
                file_content = io.BytesIO(file.read())
                file_extension = file.filename.rsplit('.', 1)[1].lower()

                # Extract text based on file type
                if file_extension == 'pdf':
                    extracted_text = extract_text_from_pdf(file_content)
                elif file_extension in ['doc', 'docx']:
                    extracted_text = extract_text_from_docx(file_content)
                else:
                    return jsonify({'status': 'error', 'message': 'Unsupported file format'}), 400

                # Parse the extracted text
                parsed_data = {
                    'user_id': str(current_user.id),  # Convert ObjectId to string
                    'personal_info': {
                        'name': extract_name(extracted_text),
                        'email': extract_email(extracted_text),
                        'phone': extract_phone(extracted_text)
                    },
                    'professional_summary': extracted_text[:500],
                    'education': extract_education(extracted_text),
                    'skills': extract_skills(extracted_text),
                    'work_experience': extract_experience(extracted_text),
                    'last_updated': datetime.utcnow().isoformat()  # Convert datetime to ISO format string
                }

                # Save to database using ObjectId for MongoDB
                db_data = parsed_data.copy()
                db_data['user_id'] = ObjectId(current_user.id)  # Convert back to ObjectId for database
                db_data['last_updated'] = datetime.utcnow()  # Use datetime object for database

                # Save to database
                result = collection_resume_details.update_one(
                    {'user_id': ObjectId(current_user.id)},
                    {'$set': db_data},
                    upsert=True
                )

                # Return JSON-serializable data
                return json.loads(json_util.dumps({
                    'status': 'success',
                    'message': 'Resume parsed successfully',
                    'data': parsed_data
                }))

            except Exception as e:
                current_app.logger.error(f'Error processing file: {str(e)}')
                return jsonify({
                    'status': 'error',
                    'message': f'Error processing file: {str(e)}'
                }), 500

        return jsonify({
            'status': 'error',
            'message': 'Invalid file type. Allowed types are PDF, DOC, and DOCX'
        }), 400

    except Exception as e:
        current_app.logger.error(f'Error in parse_resume: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@resume_maker_bp.route('/load-resume', methods=['GET'])
@login_required
def load_resume():
    try:
        # Find the user's resume in the database
        resume = collection_resume_details.find_one({'user_id': ObjectId(current_user.id)})
        
        if resume:
            # Remove MongoDB's _id field before sending
            resume.pop('_id', None)
            resume.pop('user_id', None)
            return jsonify({
                'status': 'success',
                'resume': resume
            })
        else:
            return jsonify({
                'status': 'success',
                'resume': None
            })
            
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@resume_maker_bp.route('/choose-template')
@login_required
def choose_template():
    return render_template('seeker/choose_template.html')

@resume_maker_bp.route('/generate-resume/<template_name>')
@login_required
def generate_resume(template_name):
    # Get user's resume data from database
    resume_data = collection_resume_details.find_one({'user_id': ObjectId(current_user.id)})
    
    if not resume_data:
        flash('Please create your resume first', 'error')
        return redirect(url_for('resume_maker.resume_maker'))
    
    # Remove MongoDB specific fields
    if '_id' in resume_data:
        resume_data.pop('_id')
    if 'user_id' in resume_data:
        resume_data.pop('user_id')
    
    # Ensure skills are in the correct format
    if 'skills' in resume_data:
        if isinstance(resume_data['skills'], str):
            # If skills is a string, convert it to the expected format
            resume_data['skills'] = {
                'technical': resume_data['skills'].split(','),
                'soft': []
            }
        elif isinstance(resume_data['skills'], dict):
            # Ensure each skill list contains full strings, not characters
            if 'technical' in resume_data['skills']:
                if isinstance(resume_data['skills']['technical'], str):
                    resume_data['skills']['technical'] = resume_data['skills']['technical'].split(',')
            if 'soft' in resume_data['skills']:
                if isinstance(resume_data['skills']['soft'], str):
                    resume_data['skills']['soft'] = resume_data['skills']['soft'].split(',')
    
    # Render the appropriate template
    template_map = {
        'ats': 'seeker/templates/ats_template.html',
        'bold': 'seeker/templates/bold_template.html',
        'distinct': 'seeker/templates/distinct_template.html'
    }
    
    template_path = template_map.get(template_name)
    if not template_path:
        flash('Invalid template selected', 'error')
        return redirect(url_for('resume_maker.choose_template'))
    
    return render_template(template_path, resume=resume_data)

# Add the resume parser functionality
def ats_extractor(resume_data):
    try:
        # Get API key from config file
        config_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 
                                  'config', 'ai_config.yaml')
        
        with open(config_path) as file:
            data = yaml.load(file, Loader=yaml.FullLoader)
            api_key = data.get('GEMINI_API_KEY')
        
        # Configure the Gemini API
        genai.configure(api_key=api_key)
        
        # Set up the model
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = '''
        You are an AI bot designed to act as a professional for parsing resumes. You are given with resume and your job is to extract the following information from the resume:
        
        1. full_name: The person's full name
        2. email_id: Email address
        3. github_portfolio: GitHub profile URL (if available)
        4. linkedin_id: LinkedIn profile URL (if available)
        5. summary: Look for any professional summary, profile, objective, or career statement. This might be labeled as "Summary", "Profile", "Professional Profile", "Career Objective", "About Me", etc.
        6. education: List of education entries, each with:
           - degree: The degree or certification obtained
           - institution: The school, college, or university name
           - graduation_date: When they graduated or expect to graduate
           - gpa: GPA if mentioned
        7. experience: List of work experiences, each with:
           - title: Job title
           - company: Company name
           - location: Job location (if available)
           - dates: Employment period
           - description: Bullet points or description of responsibilities and achievements
        8. skills: List of technical skills, programming languages, tools, etc.
        9. certifications: List of professional certifications (if available)
        10. projects: List of projects with descriptions (if available)
        
        Return the information in a clean JSON format. If any information is not available, use null or an empty array as appropriate.
        '''
        
        full_prompt = f"{prompt}\n\nResume content:\n{resume_data}"
        
        # Generate content
        response = model.generate_content(full_prompt)
        
        # Extract the text from the response
        data = response.text
        
        return data
    except Exception as e:
        print(f"Error generating content: {str(e)}")
        raise Exception(f"Failed to generate content: {str(e)}")

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file using pdfplumber"""
    text = ""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

# Add a new route for AI-powered resume parsing
@resume_maker_bp.route('/ai-parse-resume', methods=['POST'])
@login_required
def ai_parse_resume():
    try:
        if 'file' not in request.files:
            current_app.logger.error('No file in request')
            return jsonify({'status': 'error', 'message': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            current_app.logger.error('No filename')
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400

        # Create upload folder if it doesn't exist
        upload_path = os.path.join(current_app.root_path, 'static', 'uploads', 'resumes')
        if not os.path.exists(upload_path):
            os.makedirs(upload_path)

        # Save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(upload_path, filename)
        file.save(file_path)

        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            resume_text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith('.docx'):
            resume_text = extract_text_from_docx(file_path)
        elif filename.lower().endswith('.doc'):
            # For .doc files, you might need additional handling
            # This is a simplified approach - consider using antiword or other tools for better extraction
            try:
                # Try to open as docx first (some .doc files are actually .docx)
                resume_text = extract_text_from_docx(file_path)
            except:
                return jsonify({
                    'status': 'error',
                    'message': 'Legacy .doc format is not supported. Please convert to .docx or .pdf.'
                }), 400
        else:
            return jsonify({
                'status': 'error',
                'message': 'Unsupported file format. Please upload a PDF or DOCX file.'
            }), 400

        # Parse the resume using AI
        parsed_data = ats_extractor(resume_text)
        
        # Return the parsed data
        return jsonify({
            'status': 'success',
            'data': parsed_data
        })

    except Exception as e:
        current_app.logger.error(f'Error in AI parse_resume: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file"""
    try:
        doc = Document(docx_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
        return text
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

@resume_maker_bp.route('/save-parsed-resume', methods=['POST'])
@login_required
def save_parsed_resume():
    try:
        data = request.json
        
        resume_data = {
            'user_id': ObjectId(current_user.id),
            'personal_info': data.get('personal_info', {}),
            'professional_summary': data.get('professional_summary', ''),
            'education': data.get('education', []),
            'skills': data.get('skills', {}),
            'work_experience': data.get('work_experience', []),
            'projects': data.get('projects', []),
            'achievements': data.get('achievements', []),
            'last_updated': datetime.utcnow()
        }
        
        # Update or insert resume in database
        result = collection_resume_details.update_one(
            {'user_id': ObjectId(current_user.id)},
            {'$set': resume_data},
            upsert=True
        )
        
        return jsonify({
            'status': 'success',
            'message': 'Resume saved successfully'
        })
        
    except Exception as e:
        current_app.logger.error(f'Error saving parsed resume: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Error saving resume data: {str(e)}'
        }), 500